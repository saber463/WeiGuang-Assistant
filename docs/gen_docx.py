"""将MD计划书转换为Word文档 —— 干净重写版"""
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

MD = r'f:\java\weiguangplus\docs\项目计划书_微光同行.md'
DOCX = r'f:\java\weiguangplus\docs\项目计划书_微光同行.docx'

# ============================================================
# 工具函数
# ============================================================

def add_heading1(doc, text):
    """添加一级标题 (## 开头)"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_heading2(doc, text):
    """添加二级标题 (### 开头)"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    return p

def add_heading3(doc, text):
    """添加三级标题 (#### 开头)"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    return p

def add_para(doc, text):
    """添加普通段落，支持**加粗**"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    # 按**加粗**分割
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.bold = True
        else:
            p.add_run(part)
    return p

def add_quote(doc, text):
    """添加引用块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.size = Pt(10)
    return p

def add_code_block(doc, lines):
    """添加代码块"""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        # 设置浅灰背景
        pPr = p._element.get_or_add_pPr()
        shd = pPr.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F5F5F5',
            qn('w:val'): 'clear'
        })
        pPr.append(shd)

def add_table(doc, rows):
    """添加表格"""
    if not rows or len(rows) < 2:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci in range(num_cols):
            cell = table.rows[ri].cells[ci]
            cell_text = row[ci] if ci < len(row) else ''
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                if ri == 0:
                    for run in para.runs:
                        run.font.bold = True
                    tcPr = cell._element.get_or_add_tcPr()
                    shd = tcPr.makeelement(qn('w:shd'), {
                        qn('w:fill'): 'D6E4F0',
                        qn('w:val'): 'clear'
                    })
                    tcPr.append(shd)
    doc.add_paragraph('')  # 表后空行

# ============================================================
# 封面
# ============================================================

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('微光同行'); r.font.size = Pt(28); r.font.bold = True
r.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('智能助残服务平台商业计划书'); r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('让科技的微光，照亮每一个无声的角落')
r.font.size = Pt(14); r.font.italic = True; r.font.color.rgb = RGBColor(100,100,100)

doc.add_paragraph(''); doc.add_paragraph('')

for label, val in [
    ('项目赛道', '大学生创新创业大赛 · 红旅赛道 | 成都助残比赛 | 科大讯飞AI开发者大赛'),
    ('核心产品', 'SOS紧急求助 + 双向手语翻译 + OCR商品识别 + 家属端联动'),
    ('战略合作', '成都心融合残疾人就业基地（孙总）'),
    ('核心数据', '手势识别准确率95.9% / 端到端加密零知识安全 / 预计Y5收入2.1亿元'),
    ('版本', 'v5.0 · 2026年7月'),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{label}：{val}'); r.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 目录
# ============================================================

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('目  录'); r.font.size = Pt(20); r.font.bold = True
r.font.color.rgb = RGBColor(0, 51, 102)
doc.add_paragraph('')

# 目录生成：只扫描 MD 文件中 ## 目录 到下一个 --- 之间的内容，避免误匹配正文中的编号列表
toc_items = []
in_toc = False
for line in open(MD, encoding='utf-8').readlines():
    s = line.strip()
    if s == '## 目录':
        in_toc = True
        continue
    if in_toc and s == '---':
        break  # 目录区域结束
    if in_toc:
        m = re.match(r'^\d+\.\s+(.+)$', s)
        if m and '目录' not in m.group(1):
            toc_items.append(s)  # 保留完整行（含序号），如 "1. 项目概述"

for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.add_run(item).font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 正文：按 ## 大章节分块处理
# ============================================================

def parse_section_body(body_lines):
    """解析一个章节的body内容，返回 (paragraphs, tables, code_blocks, quotes)"""
    items = []  # 列表，每个元素是 ('type', data)
    i = 0

    while i < len(body_lines):
        line = body_lines[i]
        stripped = line.strip()

        # 代码块 ``` —— 优先级最高！必须在标题之前，防止代码块内 ###/#### 被误判为标题
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(body_lines):
                if body_lines[i].strip().startswith('```'):
                    i += 1
                    break
                code_lines.append(body_lines[i])
                i += 1
            items.append(('code', code_lines))
            continue

        # 三级标题 ####
        if stripped.startswith('#### '):
            items.append(('h3', stripped[5:].strip()))
            i += 1
            continue

        # 二级标题 ###
        if stripped.startswith('### '):
            items.append(('h2', stripped[4:].strip()))
            i += 1
            continue

        # 表格行 |
        if stripped.startswith('|') and stripped.endswith('|'):
            table_rows = []
            while i < len(body_lines):
                s = body_lines[i].strip()
                if not (s.startswith('|') and s.endswith('|')):
                    break
                cells = [c.strip() for c in s.split('|')[1:-1]]
                # 跳过分隔行
                if not all(re.match(r'^[:]?-{3,}[:]?$', c) for c in cells if c):
                    table_rows.append(cells)
                i += 1
            if table_rows:
                items.append(('table', table_rows))
            continue

        # 引用 >
        if stripped.startswith('>'):
            items.append(('quote', stripped[1:].strip()))
            i += 1
            continue

        # 分隔线或无内容
        if stripped in ('---', '') or stripped.startswith('<!--'):
            i += 1
            continue

        # 普通段落
        items.append(('para', stripped))
        i += 1

    return items


# 读取全文
with open(MD, encoding='utf-8') as f:
    all_lines = f.readlines()

# 找到第一个 ## 一、 开始的位置
start_idx = 0
for idx, line in enumerate(all_lines):
    if line.strip().startswith('## 一、'):
        start_idx = idx
        break

# 提取所有 ## 大章节
sections = []
current_title = None
current_body = []

for line in all_lines[start_idx:]:
    s = line.strip()
    if s.startswith('## ') and not s.startswith('### '):
        if current_title is not None:
            sections.append((current_title, current_body))
        current_title = s[3:].strip()
        current_body = []
    else:
        current_body.append(line.rstrip('\n'))

if current_title is not None:
    sections.append((current_title, current_body))

# 渲染每个章节
for title, body_lines in sections:
    # 去掉标题中的编号
    display_title = re.sub(r'^[一二三四五六七八九十]+、', '', title)
    add_heading1(doc, display_title)

    items = parse_section_body(body_lines)

    for item_type, data in items:
        if item_type == 'h2':
            add_heading2(doc, data)
        elif item_type == 'h3':
            add_heading3(doc, data)
        elif item_type == 'para':
            add_para(doc, data)
        elif item_type == 'quote':
            add_quote(doc, data)
        elif item_type == 'code':
            add_code_block(doc, data)
        elif item_type == 'table':
            add_table(doc, data)

# 保存
doc.save(DOCX)
print(f'Word saved: {DOCX}')
print(f'Size: {os.path.getsize(DOCX)} bytes')