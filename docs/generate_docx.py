# -*- coding: utf-8 -*-
"""
微光同行商业计划书 Markdown → Word 文档生成器
功能：读取项目计划书_微光同行.md，解析Markdown格式，生成精美的Word文档
逻辑：逐行解析MD，识别标题/表格/代码块/引用块/普通段落，分别处理
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# 配置
# ============================================================
MD_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "项目计划书_微光同行.md")
OUTPUT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "项目计划书_微光同行.docx")

# 页面设置
PAGE_WIDTH = Cm(21.0)   # A4
PAGE_HEIGHT = Cm(29.7)

# ============================================================
# 辅助函数
# ============================================================

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_paragraph(doc, text, style=None, bold=False, italic=False, font_size=None, 
                           color=None, alignment=None, font_name=None, space_after=None):
    """添加格式化段落"""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p, run

def add_rich_paragraph(doc, segments, alignment=None, space_after=None):
    """
    添加富文本段落（支持混合粗体、斜体、代码等）
    segments: list of (text, dict_of_properties)
    """
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    for text, props in segments:
        run = p.add_run(text)
        if props.get('bold'):
            run.bold = True
        if props.get('italic'):
            run.italic = True
        if props.get('code'):
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        if props.get('font_size'):
            run.font.size = Pt(props['font_size'])
        if props.get('font_name'):
            run.font.name = props['font_name']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), props['font_name'])
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def parse_inline_formatting(text):
    """解析行内格式：**粗体**、*斜体*、`代码`"""
    segments = []
    # 匹配 **粗体**、*斜体*、`代码`、普通文本
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|(.+?))'
    for match in re.finditer(pattern, text):
        if match.group(2) is not None:
            segments.append((match.group(2), {'bold': True}))
        elif match.group(3) is not None:
            segments.append((match.group(3), {'italic': True}))
        elif match.group(4) is not None:
            segments.append((match.group(4), {'code': True}))
        elif match.group(5) is not None:
            segments.append((match.group(5), {}))
    return segments if segments else [(text, {})]

def parse_table_row(line):
    """解析表格行，返回单元格列表"""
    # 去掉首尾的 |
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [cell.strip() for cell in line.split('|')]

def is_separator_row(cells):
    """判断是否为表格分隔行（如 |---|---|）"""
    return all(re.match(r'^[-:]+$', c.strip()) for c in cells)

def add_table_to_doc(doc, rows_data, col_widths=None):
    """添加表格到文档，带格式"""
    if not rows_data:
        return
    num_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 遍历所有行
    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = row.cells[j]
            # 清除默认段落
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            
            # 表头行（第一行）加粗+浅灰背景
            if i == 0:
                run.bold = True
                run.font.size = Pt(9.5)
                set_cell_shading(cell, 'D9E2F3')  # 浅蓝色背景
    
    # 添加表后空行
    doc.add_paragraph()
    return table

def add_code_block(doc, code_lines):
    """添加代码块"""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(51, 51, 51)
    doc.add_paragraph()  # 代码块后空行

def add_quote_block(doc, quote_lines):
    """添加引用块"""
    for line in quote_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1)
        # 左侧灰色竖线效果（用缩进模拟）
        run = p.add_run(line)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ============================================================
# 主解析器
# ============================================================

def generate_docx(md_path, output_path):
    """主函数：解析Markdown并生成Word文档"""
    doc = Document()
    
    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    # ========== 设置默认字体 ==========
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ========== 读取MD文件 ==========
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # ========== 解析状态机 ==========
    i = 0
    in_cover = False
    in_toc = False
    in_code_block = False
    in_table = False
    in_quote = False
    code_lines = []
    table_rows = []
    quote_lines = []
    cover_ended = False
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # ========== 代码块处理 ==========
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # 开始代码块
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # ========== 引用块处理 ==========
        if line.strip().startswith('> '):
            quote_lines.append(line.strip()[2:])
            in_quote = True
            i += 1
            continue
        
        if in_quote:
            # 引用块结束
            add_quote_block(doc, quote_lines)
            quote_lines = []
            in_quote = False
            # 不增加i，继续处理当前行
            continue
        
        # ========== 表格处理 ==========
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
            row_cells = parse_table_row(line)
            # 跳过纯分隔符行
            if not is_separator_row(row_cells):
                table_rows.append(row_cells)
            i += 1
            continue
        
        if in_table:
            # 表格结束
            add_table_to_doc(doc, table_rows)
            table_rows = []
            in_table = False
            # 不增加i，继续处理当前行
            continue
        
        # ========== 空行跳过 ==========
        if line.strip() == '':
            i += 1
            continue
        
        # ========== 水平线 ==========
        if line.strip() == '---':
            # 封面结束标记
            if not cover_ended and any('封面' in t for t in [l.strip() for l in lines[max(0,i-10):i]]):
                cover_ended = True
                doc.add_page_break()
            # 目录结束标记
            elif any('目录' in t for t in [l.strip() for l in lines[max(0,i-5):i]]):
                doc.add_page_break()
            i += 1
            continue
        
        # ========== 标题处理 ==========
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()
            
            # 处理标题中的行内格式
            heading = doc.add_heading(level=level)
            run = heading.add_run(title_text)
            
            if level == 1:
                run.font.size = Pt(22)
            elif level == 2:
                run.font.size = Pt(16)
            elif level == 3:
                run.font.size = Pt(13)
            else:
                run.font.size = Pt(11)
            
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue
        
        # ========== 普通段落 ==========
        stripped = line.strip()
        if stripped:
            # 特殊处理：场景标题（以**开头的加粗标题）
            if stripped.startswith('**场景') or stripped.startswith('**四个场景'):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                segments = parse_inline_formatting(stripped)
                for text, props in segments:
                    run = p.add_run(text)
                    if props.get('bold'):
                        run.bold = True
                    run.font.size = Pt(11)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                segments = parse_inline_formatting(stripped)
                for text, props in segments:
                    run = p.add_run(text)
                    if props.get('bold'):
                        run.bold = True
                    if props.get('italic'):
                        run.italic = True
                    if props.get('code'):
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10.5)
        
        i += 1
    
    # ========== 保存文档 ==========
    doc.save(output_path)
    print(f"Word文档已生成: {output_path}")
    return output_path

# ============================================================
# 入口
# ============================================================
if __name__ == '__main__':
    generate_docx(MD_FILE, OUTPUT_FILE)