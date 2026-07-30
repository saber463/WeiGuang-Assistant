# -*- coding: utf-8 -*-
"""
Markdown 转 Word 文档转换脚本
功能：将 项目计划书_微光同行.md 转换为格式化的 .docx 文件
逻辑：状态机逐行解析 Markdown，识别标题/表格/代码块/引用块/数学公式等元素，
      并应用对应的 Word 格式（字体、颜色、背景、缩进等）
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

# ============================================================
# 常量定义
# ============================================================
# 颜色
DEEP_BLUE = RGBColor(0, 51, 102)       # 一级标题深蓝色
LIGHT_BLUE_BG = "D6E4F0"               # 表格表头背景色
LIGHT_GRAY_BG = "F5F5F5"               # 代码块背景色
QUOTE_GRAY = RGBColor(102, 102, 102)    # 引用块文字颜色

# 字体
BODY_FONT = "微软雅黑"
CODE_FONT = "Consolas"

# 字号
H1_SIZE = Pt(22)
H2_SIZE = Pt(16)
H3_SIZE = Pt(13)
H4_SIZE = Pt(11)
BODY_SIZE = Pt(10.5)
TABLE_SIZE = Pt(9)
CODE_SIZE = Pt(9)
COVER_TITLE_SIZE = Pt(28)
COVER_SUBTITLE_SIZE = Pt(14)
COVER_INFO_SIZE = Pt(12)
TOC_SIZE = Pt(12)

# 间距
PARA_SPACE_AFTER = Pt(4)
HEADING_SPACE_BEFORE = Pt(12)
HEADING_SPACE_AFTER = Pt(6)


def set_font(run, font_name, size, bold=False, italic=False, color=None):
    """
    设置 run 的字体属性（同时设置西文和东亚字体，确保中文正确渲染）
    :param run: python-docx Run 对象
    :param font_name: 字体名称
    :param size: 字号 Pt 对象
    :param bold: 是否加粗
    :param italic: 是否斜体
    :param color: RGBColor 颜色对象
    """
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, space_before=0, space_after=PARA_SPACE_AFTER, line_spacing=1.15):
    """
    设置段落间距
    :param paragraph: python-docx Paragraph 对象
    :param space_before: 段前间距 Pt
    :param space_after: 段后间距 Pt
    :param line_spacing: 行距倍数
    """
    pf = paragraph.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing


def set_cell_shading(cell, color_hex):
    """
    设置表格单元格背景色
    :param cell: python-docx Cell 对象
    :param color_hex: 六位十六进制颜色码（不含#）
    """
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._element.tcPr.append(shading_elm)


def set_paragraph_shading(paragraph, color_hex):
    """
    设置段落背景色
    :param paragraph: python-docx Paragraph 对象
    :param color_hex: 六位十六进制颜色码（不含#）
    """
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    paragraph._element.pPr.append(shading_elm)


def add_horizontal_line(doc):
    """在文档中添加水平分隔线（使用底部边框模拟）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.pPr
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_empty_line(doc, height_pt=6):
    """添加一个空行作为间距"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(height_pt)


def add_body_paragraph(doc, text, bold=False, italic=False):
    """
    添加正文段落
    :param doc: Document 对象
    :param text: 段落文本
    :param bold: 是否加粗
    :param italic: 是否斜体
    :return: Paragraph 对象
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    run = p.add_run(text)
    set_font(run, BODY_FONT, BODY_SIZE, bold=bold, italic=italic)
    return p


def add_heading_paragraph(doc, text, level):
    """
    添加标题段落
    :param doc: Document 对象
    :param text: 标题文本
    :param level: 标题级别 1-4
    :return: Paragraph 对象
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if level == 1:
        set_paragraph_spacing(p, space_before=HEADING_SPACE_BEFORE, space_after=HEADING_SPACE_AFTER, line_spacing=1.2)
        run = p.add_run(text)
        set_font(run, BODY_FONT, H1_SIZE, bold=True, color=DEEP_BLUE)
    elif level == 2:
        set_paragraph_spacing(p, space_before=HEADING_SPACE_BEFORE, space_after=HEADING_SPACE_AFTER, line_spacing=1.2)
        run = p.add_run(text)
        set_font(run, BODY_FONT, H2_SIZE, bold=True)
    elif level == 3:
        set_paragraph_spacing(p, space_before=Pt(8), space_after=Pt(4), line_spacing=1.2)
        run = p.add_run(text)
        set_font(run, BODY_FONT, H3_SIZE, bold=True)
    elif level == 4:
        set_paragraph_spacing(p, space_before=Pt(6), space_after=Pt(3), line_spacing=1.2)
        run = p.add_run(text)
        set_font(run, BODY_FONT, H4_SIZE, bold=True)
    return p


def add_quote_paragraph(doc, lines):
    """
    添加引用块段落（斜体，灰色，左缩进）
    :param doc: Document 对象
    :param lines: 引用块文本行列表
    """
    for line in lines:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=Pt(0), space_after=Pt(2), line_spacing=1.15)
        # 左缩进
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(line)
        set_font(run, BODY_FONT, BODY_SIZE, italic=True, color=QUOTE_GRAY)


def add_code_block(doc, lines):
    """
    添加代码块（Consolas等宽字体，浅灰背景，保留原样）
    :param doc: Document 对象
    :param lines: 代码块文本行列表
    """
    for line in lines:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=Pt(0), space_after=Pt(0), line_spacing=1.0)
        set_paragraph_shading(p, LIGHT_GRAY_BG)
        # 代码块左缩进
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line if line else " ")
        set_font(run, CODE_FONT, CODE_SIZE)


def add_table_from_data(doc, rows_data, has_header=True):
    """
    根据二维数组数据创建表格
    :param doc: Document 对象
    :param rows_data: 二维数组，每行是一个列表
    :param has_header: 第一行是否为表头
    :return: Table 对象
    """
    if not rows_data:
        return None

    num_cols = max(len(row) for row in rows_data)
    # 补齐不足的列
    for row in rows_data:
        while len(row) < num_cols:
            row.append("")

    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            # 清除默认段落
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text.strip())
            set_font(run, BODY_FONT, TABLE_SIZE, bold=(has_header and i == 0))

            # 表头行设置背景色
            if has_header and i == 0:
                set_cell_shading(cell, LIGHT_BLUE_BG)

            # 设置单元格段落间距
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0

    return table


def parse_markdown_table(lines, start_idx):
    """
    解析Markdown表格，返回 (rows_data, end_idx)
    :param lines: 所有行列表
    :param start_idx: 表格起始行索引
    :return: (二维数组rows_data, 表格结束后的下一行索引)
    """
    rows_data = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line or not line.startswith('|'):
            break
        # 跳过分隔行（如 |------|------|）
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        # 解析单元格
        cells = [c.strip() for c in line.split('|')]
        # 去掉首尾空元素
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        rows_data.append(cells)
        i += 1
    return rows_data, i


def parse_blockquote(lines, start_idx):
    """
    解析引用块，返回 (文本行列表, end_idx)
    :param lines: 所有行列表
    :param start_idx: 引用块起始行索引
    :return: (文本行列表, 引用块结束后的下一行索引)
    """
    quote_lines = []
    i = start_idx
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped or not line.startswith('>'):
            break
        # 去掉开头的 > 和空格
        content = re.sub(r'^>\s?', '', line)
        quote_lines.append(content)
        i += 1
    return quote_lines, i


def parse_code_block(lines, start_idx):
    """
    解析代码块，返回 (代码行列表, end_idx)
    :param lines: 所有行列表
    :param start_idx: 代码块起始行索引（```所在行）
    :return: (代码行列表, 代码块结束后的下一行索引)
    """
    code_lines = []
    i = start_idx + 1  # 跳过开头的 ```
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```'):
            i += 1
            break
        code_lines.append(line)
        i += 1
    return code_lines, i


def extract_cover_data(lines):
    """
    从Markdown中提取封面数据
    :param lines: 所有行列表
    :return: dict 封面信息
    """
    cover = {}
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped == '## 封面':
            # 找到封面表格
            for j in range(i + 1, min(i + 30, len(lines))):
                if lines[j].strip().startswith('|'):
                    rows, _ = parse_markdown_table(lines, j)
                    for row in rows:
                        if len(row) >= 2:
                            cover[row[0].strip()] = row[1].strip()
                    break
            break
    return cover


def extract_toc_lines(lines):
    """
    从Markdown中提取目录行
    :param lines: 所有行列表
    :return: 目录文本行列表
    """
    toc_lines = []
    in_toc = False
    for line in lines:
        stripped = line.strip()
        if stripped == '## 目录':
            in_toc = True
            continue
        if in_toc:
            if stripped == '---' or stripped.startswith('## '):
                break
            if stripped:
                # 去掉编号前的数字
                toc_lines.append(stripped)
    return toc_lines


def create_cover_page(doc, cover_data, title_line, meta_lines):
    """
    创建封面页
    :param doc: Document 对象
    :param cover_data: 封面信息字典
    :param title_line: 主标题行
    :param meta_lines: 元信息行（引用块内容）
    """
    # 添加一些空行使封面内容居中
    for _ in range(6):
        add_empty_line(doc, height_pt=14)

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, space_before=Pt(0), space_after=Pt(8), line_spacing=1.3)
    run = p.add_run(title_line.replace('# ', '').strip())
    set_font(run, BODY_FONT, COVER_TITLE_SIZE, bold=True, color=DEEP_BLUE)

    # 副标题 / Slogan
    if '项目Slogan' in cover_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, space_before=Pt(0), space_after=Pt(20), line_spacing=1.3)
        run = p.add_run(cover_data['项目Slogan'])
        set_font(run, BODY_FONT, COVER_SUBTITLE_SIZE, italic=True, color=QUOTE_GRAY)

    # 分隔线
    add_horizontal_line(doc)
    add_empty_line(doc, height_pt=10)

    # 封面信息
    cover_keys = ['项目名称', '参赛赛道', '核心产品', '战略合作', '核心数据', '团队', '日期']
    for key in cover_keys:
        if key in cover_data:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, space_before=Pt(0), space_after=Pt(4), line_spacing=1.3)
            run = p.add_run(f"{key}：{cover_data[key]}")
            set_font(run, BODY_FONT, COVER_INFO_SIZE)

    # 分页
    doc.add_page_break()


def create_toc_page(doc, toc_lines):
    """
    创建目录页
    :param doc: Document 对象
    :param toc_lines: 目录文本行列表
    """
    # 目录标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, space_before=Pt(0), space_after=Pt(16), line_spacing=1.3)
    run = p.add_run("目  录")
    set_font(run, BODY_FONT, H1_SIZE, bold=True)

    add_horizontal_line(doc)
    add_empty_line(doc, height_pt=8)

    # 目录条目
    for i, line in enumerate(toc_lines):
        # 清理目录行（去掉编号前缀如 "1. "）
        cleaned = re.sub(r'^\d+\.\s*', '', line)
        p = doc.add_paragraph()
        set_paragraph_spacing(p, space_before=Pt(0), space_after=Pt(3), line_spacing=1.5)
        run = p.add_run(f"{cleaned}")
        set_font(run, BODY_FONT, TOC_SIZE)

    # 分页
    doc.add_page_break()


def process_inline_formatting(text):
    """
    处理Markdown行内格式：**加粗**、*斜体*、~~删除线~~、`代码`
    返回 (格式化后的纯文本, 格式标记列表)
    由于python-docx的run级别格式更灵活，这里简化为返回纯文本
    实际格式化在add_body_paragraph_with_inline中处理
    """
    return text


def add_formatted_paragraph(doc, text):
    """
    添加带行内格式的段落（处理**加粗**、*斜体*等）
    :param doc: Document 对象
    :param text: 段落文本
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p)

    # 使用正则分割文本，识别 **...** 和 *...*
    # 模式：**text** 或 *text*（但不匹配 ** 内部的 *）
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|[^*]+)'
    parts = re.findall(pattern, text)

    for part in parts:
        full_match = part[0]
        bold_content = part[1]
        italic_content = part[2]

        if bold_content:
            run = p.add_run(bold_content)
            set_font(run, BODY_FONT, BODY_SIZE, bold=True)
        elif italic_content:
            run = p.add_run(italic_content)
            set_font(run, BODY_FONT, BODY_SIZE, italic=True)
        else:
            # 普通文本
            run = p.add_run(full_match)
            set_font(run, BODY_FONT, BODY_SIZE)

    return p


def convert_markdown_to_docx(md_path, docx_path):
    """
    主转换函数：将Markdown文件转换为Word文档
    :param md_path: Markdown文件路径
    :param docx_path: 输出Word文档路径
    """
    # 读取Markdown文件
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    # 提取封面数据和目录数据
    cover_data = extract_cover_data(lines)
    toc_lines = extract_toc_lines(lines)

    # 提取主标题和元信息
    title_line = ""
    meta_lines = []
    for line in lines[:10]:
        stripped = line.strip()
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title_line = stripped
        elif line.startswith('>'):
            meta_lines.append(re.sub(r'^>\s?', '', line))

    # 创建Word文档
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = PARA_SPACE_AFTER
    style.paragraph_format.line_spacing = 1.15

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========================================
    # 封面页
    # ========================================
    create_cover_page(doc, cover_data, title_line, meta_lines)

    # ========================================
    # 目录页
    # ========================================
    create_toc_page(doc, toc_lines)

    # ========================================
    # 正文内容
    # 跳过封面和目录部分，从 "## 一、项目概述" 开始
    # ========================================
    body_start = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped == '## 一、项目概述':
            body_start = i
            break

    # 状态机变量
    in_code_block = False
    in_table = False
    in_blockquote = False
    code_block_lines = []
    table_start_idx = 0
    blockquote_lines = []

    i = body_start
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ---- 处理代码块 ----
        if stripped.startswith('```'):
            if in_code_block:
                # 代码块结束
                add_code_block(doc, code_block_lines)
                code_block_lines = []
                in_code_block = False
                i += 1
                continue
            else:
                # 代码块开始
                in_code_block = True
                code_block_lines = []
                i += 1
                continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # ---- 处理表格 ----
        if stripped.startswith('|') and not in_blockquote:
            if not in_table:
                in_table = True
                table_start_idx = i
            i += 1
            continue

        if in_table:
            # 表格结束，解析并输出
            rows_data, _ = parse_markdown_table(lines, table_start_idx)
            if rows_data:
                add_empty_line(doc, height_pt=4)
                add_table_from_data(doc, rows_data, has_header=True)
                add_empty_line(doc, height_pt=4)
            in_table = False
            # 不递增 i，因为当前行不是表格内容，需要重新处理
            continue

        # ---- 处理引用块 ----
        if line.startswith('>') and not in_blockquote:
            in_blockquote = True
            blockquote_lines = []
            blockquote_lines.append(re.sub(r'^>\s?', '', line))
            i += 1
            continue

        if in_blockquote:
            if line.startswith('>'):
                blockquote_lines.append(re.sub(r'^>\s?', '', line))
                i += 1
                continue
            else:
                # 引用块结束
                add_quote_paragraph(doc, blockquote_lines)
                blockquote_lines = []
                in_blockquote = False
                # 继续处理当前行
                continue

        # ---- 处理分隔线 ----
        if stripped == '---' or stripped == '***':
            add_empty_line(doc, height_pt=2)
            add_horizontal_line(doc)
            add_empty_line(doc, height_pt=2)
            i += 1
            continue

        # ---- 处理标题 ----
        if stripped.startswith('#'):
            # 计算标题级别
            level = 0
            for ch in stripped:
                if ch == '#':
                    level += 1
                else:
                    break
            heading_text = stripped[level:].strip()
            if 1 <= level <= 4:
                if level == 1:
                    add_empty_line(doc, height_pt=6)
                add_heading_paragraph(doc, heading_text, level)
            i += 1
            continue

        # ---- 处理数学公式 $$...$$ ----
        if stripped.startswith('$$') and stripped.endswith('$$'):
            formula = stripped[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, space_before=Pt(4), space_after=Pt(4), line_spacing=1.3)
            run = p.add_run(formula)
            set_font(run, CODE_FONT, CODE_SIZE)
            i += 1
            continue

        # ---- 处理空行 ----
        if not stripped:
            i += 1
            continue

        # ---- 处理普通段落 ----
        add_formatted_paragraph(doc, stripped)
        i += 1

    # 处理末尾未闭合的元素
    if in_code_block and code_block_lines:
        add_code_block(doc, code_block_lines)
    if in_blockquote and blockquote_lines:
        add_quote_paragraph(doc, blockquote_lines)

    # 保存文档
    doc.save(docx_path)
    print(f"✅ Word文档已生成：{docx_path}")


if __name__ == '__main__':
    md_file = os.path.join(os.path.dirname(__file__), '项目计划书_微光同行.md')
    docx_file = os.path.join(os.path.dirname(__file__), '项目计划书_微光同行.docx')

    if not os.path.exists(md_file):
        print(f"❌ 源文件不存在：{md_file}")
        exit(1)

    print(f"📄 源文件：{md_file}")
    print(f"📝 目标文件：{docx_file}")
    print("🔄 正在转换...")

    convert_markdown_to_docx(md_file, docx_file)

    # 验证
    if os.path.exists(docx_file):
        size_kb = os.path.getsize(docx_file) / 1024
        print(f"✅ 验证通过！文件大小：{size_kb:.1f} KB")
    else:
        print("❌ 生成失败！")
        exit(1)